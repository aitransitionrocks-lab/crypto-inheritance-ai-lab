"""Report Agent – generates 5 separate strategic reports plus a combined .docx."""

from __future__ import annotations

import logging
from pathlib import Path
from datetime import datetime, timezone

from docx import Document as DocxDocument
from docx.shared import Pt, Inches

from agents.base import BaseAgent
from config.settings import settings

logger = logging.getLogger(__name__)


class ReportAgent(BaseAgent):
    name = "report"
    role = "Strategic Report Compiler"
    system_prompt = (
        "You are an expert business strategist and report writer for LegacyGuard, a crypto "
        "inheritance protocol and security infrastructure layer. You compile research findings "
        "into sharp, decision-useful strategic documents. Your writing is direct, specific, "
        "and avoids fluff. Every report includes explicit recommendations, scorecards, "
        "comparison tables, go/no-go conclusions, key assumptions, and open questions. "
        "You write for founders and investors who need to make real decisions."
    )

    async def run(self, startup_idea: str) -> dict[str, str]:
        # Collect all section outputs from memory
        sections_queries: list[tuple[str, str]] = [
            ("Problem Analysis", "problem analysis crypto inheritance segments painkiller"),
            ("Market Opportunity", "market opportunity TAM SAM SOM segments B2B"),
            ("Competitor Landscape", "competitor landscape whitespace categories"),
            ("Product Design", "product design MVP variants API SDK infrastructure"),
            ("Security Architecture", "security architecture cryptographic custodian options"),
            ("Legal Strategy", "legal strategy provider classification jurisdiction custody"),
            ("Business Model", "business model revenue strategies scorecard unit economics"),
            ("Go-To-Market Strategy", "GTM partnership pathways infrastructure pilot"),
        ]

        collected: dict[str, str] = {}
        for title, query in sections_queries:
            docs = self.memory.query(query, n_results=2, where={"type": "output"})
            collected[title] = "\n\n".join(d["document"] for d in docs) if docs else ""

        all_content: str = "\n\n".join(
            f"## {k}\n{v}" for k, v in collected.items() if v
        )

        # ---- Report 1: Executive Strategy Report ----
        report_1: str = await self.think(
            prompt=(
                "Write Report 1: 'LegacyGuard — Protocol + Infrastructure Strategy'\n\n"
                "This is the executive strategy report. Structure:\n\n"
                "# LegacyGuard — Protocol + Infrastructure Strategy\n\n"
                "## Executive Summary (500 words max)\n"
                "The problem, the opportunity, why protocol + infrastructure (not just app).\n\n"
                "## Strategic Positioning\n"
                "Why LegacyGuard should position as infrastructure, not consumer SaaS.\n"
                "Comparison table: Consumer SaaS vs Infrastructure vs Protocol positioning.\n\n"
                "## Market Opportunity Summary\n"
                "Combined TAM/SAM/SOM table across all segments.\n\n"
                "## Competitive Advantage\n"
                "What LegacyGuard can own that competitors cannot.\n"
                "Defensibility assessment (1-10 scale across 5 moat types).\n\n"
                "## Architecture Decision\n"
                "Recommended technical architecture and why.\n\n"
                "## Business Model Recommendation\n"
                "Primary + secondary revenue model. Unit economics summary.\n\n"
                "## Go/No-Go Scorecard\n"
                "| Criteria | Score (1-10) | Evidence | Risk |\n"
                "Cover: Market size, timing, defensibility, execution complexity, "
                "regulatory risk, team-market fit, funding feasibility.\n\n"
                "## Key Assumptions (Must Validate)\n"
                "Top 5 assumptions that must be true. How to test each.\n\n"
                "## Open Questions\n"
                "5-10 questions that remain unanswered and need founder investigation.\n\n"
                "## Conclusion: Build or Don't Build?\n"
                "Clear recommendation with conditions.\n\n"
                f"Startup Idea: {startup_idea}"
            ),
            context=all_content,
        )
        self.store_output("report_executive_strategy", report_1)

        # ---- Report 2: VC Investment Memo ----
        report_2: str = await self.think(
            prompt=(
                "Write Report 2: 'Should LegacyGuard Be Built as a Venture-Scale Company?'\n\n"
                "This is a VC investment memo format. Structure:\n\n"
                "# VC Investment Memo: LegacyGuard\n\n"
                "## Thesis (1 paragraph)\n"
                "Why this could be a venture-scale opportunity.\n\n"
                "## The Problem (with data)\n"
                "Quantified problem. Why existing solutions fail.\n\n"
                "## The Solution\n"
                "What LegacyGuard builds. Key differentiator.\n\n"
                "## Market Size\n"
                "TAM/SAM/SOM with methodology. Comparison to funded analogues.\n\n"
                "## Business Model & Unit Economics\n"
                "Revenue model, pricing, LTV:CAC, path to $10M ARR.\n\n"
                "## Competitive Landscape\n"
                "Who else is doing this? Why LegacyGuard wins.\n\n"
                "## Go-To-Market\n"
                "How the first $1M in revenue gets generated.\n\n"
                "## Risks & Mitigations\n"
                "| Risk | Likelihood | Impact | Mitigation |\n\n"
                "## Financial Projections (3 years)\n"
                "| Metric | Year 1 | Year 2 | Year 3 |\n"
                "Revenue, users/customers, burn rate, team size.\n\n"
                "## Funding Ask\n"
                "Pre-seed/seed amount, use of funds, milestones.\n\n"
                "## Venture Scale Assessment\n"
                "Can this be a $100M+ revenue company? Under what conditions?\n"
                "Score on: Market size, winner-take-most dynamics, margin profile, "
                "network effects, retention.\n\n"
                "## Conclusion: Invest or Pass?\n"
                "Clear recommendation for a hypothetical seed investor.\n\n"
                f"Startup Idea: {startup_idea}"
            ),
            context=all_content,
        )
        self.store_output("report_vc_memo", report_2)

        # ---- Report 3: MVP Recommendation ----
        report_3: str = await self.think(
            prompt=(
                "Write Report 3: 'What Should LegacyGuard Build First?'\n\n"
                "This is the MVP recommendation report. Structure:\n\n"
                "# MVP Recommendation: What to Build First\n\n"
                "## The 3 MVP Options (from Product Design)\n"
                "Summarize MVP-A (fastest), MVP-B (B2B), MVP-C (protocol).\n\n"
                "## Scoring Matrix\n"
                "| Criteria | Weight | MVP-A | MVP-B | MVP-C |\n"
                "Criteria: Speed to market, cost to build, revenue potential, "
                "learning value, defensibility, team feasibility, investor appeal.\n"
                "Weighted total score for each.\n\n"
                "## Recommended MVP\n"
                "Which one and why. Be specific.\n\n"
                "## MVP Specification\n"
                "- Exact feature list (what's in, what's out)\n"
                "- Target user/customer\n"
                "- Success metrics (what proves the MVP works)\n"
                "- Timeline (week-by-week for 12 weeks)\n"
                "- Team needed (roles, count)\n"
                "- Budget estimate\n\n"
                "## Validation Plan\n"
                "How to test the 5 riskiest assumptions before building.\n"
                "| Assumption | Test Method | Success Criteria | Timeline |\n\n"
                "## Kill Criteria\n"
                "What signals should make the founder STOP and pivot?\n\n"
                "## After MVP: What's Next?\n"
                "If MVP succeeds, what's the next 6-month plan?\n\n"
                f"Startup Idea: {startup_idea}"
            ),
            context=all_content,
        )
        self.store_output("report_mvp_recommendation", report_3)

        # ---- Report 4: Partnership Playbook ----
        report_4: str = await self.think(
            prompt=(
                "Write Report 4: 'How LegacyGuard Can Enter the Market Through "
                "Infrastructure Partnerships'\n\n"
                "This is the partnership playbook. Structure:\n\n"
                "# Partnership Playbook: Infrastructure-First Market Entry\n\n"
                "## Why Partnerships First?\n"
                "The case for B2B distribution over direct consumer acquisition.\n\n"
                "## Top 5 Partnership Pathways (Detailed)\n"
                "For EACH pathway:\n"
                "### Pathway: [Partner Type]\n"
                "- Specific target companies (name 3-5)\n"
                "- Decision maker title and typical buying process\n"
                "- Value proposition (their language, their pain)\n"
                "- Integration scope and timeline\n"
                "- Commercial model (pricing, revenue share)\n"
                "- Pilot offer\n"
                "- Outreach strategy (warm intros, conferences, direct)\n"
                "- Expected close timeline\n\n"
                "## Partnership Priority Matrix\n"
                "| Partner Type | Impact | Feasibility | Speed | Rev Potential | "
                "Strategic Value | Priority |\n\n"
                "## Outreach Templates\n"
                "Draft 3 outreach messages (cold email, warm intro, conference follow-up) "
                "tailored to the #1 priority partner type.\n\n"
                "## Pilot Program Structure\n"
                "- Duration, scope, success criteria\n"
                "- What LegacyGuard provides free vs paid\n"
                "- Data/metrics to collect during pilot\n\n"
                "## Common Objections & Responses\n"
                "| Objection | Response | Evidence |\n\n"
                "## Partnership Funnel Metrics\n"
                "Expected conversion rates at each stage.\n\n"
                f"Startup Idea: {startup_idea}"
            ),
            context=all_content,
        )
        self.store_output("report_partnership_playbook", report_4)

        # ---- Report 5: Founder Action Plan ----
        report_5: str = await self.think(
            prompt=(
                "Write Report 5: '30-Day, 60-Day, 90-Day Founder Action Plan'\n\n"
                "This is the concrete action plan. Structure:\n\n"
                "# Founder Action Plan: 30 / 60 / 90 Days\n\n"
                "## Pre-Requisites (Week 0)\n"
                "What the founder should have in place before starting the clock.\n\n"
                "## Days 1-30: Validate & Prepare\n"
                "Week-by-week breakdown:\n"
                "### Week 1: Problem Validation\n"
                "- Specific actions (interviews, surveys, data gathering)\n"
                "- Target: N interviews with which personas\n"
                "- Key questions to ask\n"
                "### Week 2: Solution Hypothesis\n"
                "- Refine architecture choice based on validation\n"
                "- Draft technical spec\n"
                "### Week 3: Market Validation\n"
                "- Partner conversations (which companies, what to pitch)\n"
                "- Competitive positioning refinement\n"
                "### Week 4: Go/No-Go Decision\n"
                "- Decision framework: build, pivot, or stop\n"
                "- Criteria for each outcome\n\n"
                "## Days 31-60: Build MVP\n"
                "Week-by-week breakdown:\n"
                "### Week 5-6: Core Build\n"
                "### Week 7-8: Integration & Testing\n"
                "Specific deliverables each week.\n\n"
                "## Days 61-90: Launch & Learn\n"
                "Week-by-week breakdown:\n"
                "### Week 9-10: Pilot Launch\n"
                "### Week 11-12: Iterate & Expand\n"
                "Specific deliverables each week.\n\n"
                "## Key Milestones Checklist\n"
                "| Milestone | Target Date | Success Metric | Status |\n\n"
                "## Budget for 90 Days\n"
                "| Category | Monthly Cost | 90-Day Total |\n\n"
                "## Team Needs\n"
                "What roles are needed, when to hire, where to find them.\n\n"
                "## Risk Register\n"
                "| Risk | Probability | Impact | Mitigation | Owner |\n\n"
                "## Decision Log Template\n"
                "Key decisions the founder will need to make, with decision criteria.\n\n"
                f"Startup Idea: {startup_idea}"
            ),
            context=all_content,
        )
        self.store_output("report_founder_action_plan", report_5)

        # ---- Assemble all reports ----
        report_sections: dict[str, str] = {
            "Executive Strategy Report": report_1,
            "VC Investment Memo": report_2,
            "MVP Recommendation": report_3,
            "Partnership Playbook": report_4,
            "Founder Action Plan": report_5,
        }

        # Write individual markdown files
        for slug, content in [
            ("01_executive_strategy", report_1),
            ("02_vc_investment_memo", report_2),
            ("03_mvp_recommendation", report_3),
            ("04_partnership_playbook", report_4),
            ("05_founder_action_plan", report_5),
        ]:
            md_path = self._write_markdown_report(slug, startup_idea, content)
            logger.info("Report written: %s", md_path)

        # Write combined docx
        docx_path = self._write_combined_docx(startup_idea, report_sections)
        logger.info("Combined report written: %s", docx_path)

        return report_sections

    # ----------------------------------------------------------- markdown
    def _write_markdown_report(
        self, slug: str, idea: str, content: str
    ) -> Path:
        settings.reports_dir.mkdir(parents=True, exist_ok=True)
        now = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        header = (
            f"# LegacyGuard: {slug.replace('_', ' ').title()}\n"
            f"*Startup Idea: {idea}*\n"
            f"*Generated {now} by LegacyGuard AI Strategy Lab*\n\n"
            "---\n\n"
        )
        full_content = header + (content.strip() if content else "*No data collected.*")
        path = settings.reports_dir / f"legacyguard_{slug}.md"
        path.write_text(full_content, encoding="utf-8")
        return path

    # ----------------------------------------------------------- docx
    def _write_combined_docx(
        self, idea: str, reports: dict[str, str]
    ) -> Path:
        settings.reports_dir.mkdir(parents=True, exist_ok=True)
        doc = DocxDocument()

        # Title page
        doc.add_heading("LegacyGuard: Strategic Analysis", level=0)
        doc.add_heading("Crypto Inheritance Protocol & Infrastructure Layer", level=1)
        now = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        doc.add_paragraph(f"Startup Idea: {idea}")
        doc.add_paragraph(f"Generated {now} by LegacyGuard AI Strategy Lab")
        doc.add_paragraph("")
        doc.add_paragraph("This document contains 5 strategic reports:")
        for i, title in enumerate(reports.keys(), 1):
            doc.add_paragraph(f"  {i}. {title}", style="List Number")
        doc.add_page_break()

        # Each report as a section
        for title, content in reports.items():
            doc.add_heading(title, level=1)
            doc.add_paragraph("")

            for paragraph in (content or "No data collected.").split("\n\n"):
                cleaned = paragraph.strip()
                if not cleaned:
                    continue
                if cleaned.startswith("#### "):
                    doc.add_heading(cleaned[5:], level=4)
                elif cleaned.startswith("### "):
                    doc.add_heading(cleaned[4:], level=3)
                elif cleaned.startswith("## "):
                    doc.add_heading(cleaned[3:], level=2)
                elif cleaned.startswith("# "):
                    doc.add_heading(cleaned[2:], level=1)
                elif cleaned.startswith("| "):
                    # Table content — render as preformatted
                    p = doc.add_paragraph()
                    run = p.add_run(cleaned)
                    run.font.size = Pt(9)
                    run.font.name = "Courier New"
                else:
                    doc.add_paragraph(cleaned)

            doc.add_page_break()

        path = settings.reports_dir / "legacyguard_combined_strategy.docx"
        doc.save(str(path))
        return path
